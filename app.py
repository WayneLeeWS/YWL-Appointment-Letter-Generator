import io
import os
import re
import subprocess
import tempfile
import datetime
import requests
import streamlit as st
import streamlit.components.v1 as components
from docx import Document

# ==========================================
# PAGE CONFIGURATION & CSS CUSTOMIZATION
# ==========================================
hide_streamlit_style = """
    <style>
    /* Hides the top-right menu and GitHub icon */
    [data-testid="stToolbar"] {
        visibility: hidden !important;
    }
    
    /* Hides any anchor tag pointing to GitHub */
    a[href^="https://github.com"] {
        display: none !important;
    }
    
    /* Hides the Streamlit footer ("Made with Streamlit") */
    footer {
        visibility: hidden !important;
    }
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

st.set_page_config(
    page_title="Document Generator | YWL Holding",
    page_icon="📄",
    layout="wide"
)

# Hide Streamlit's "Press Enter to submit" helper text in form inputs
st.markdown("""
    <style>
    div[data-testid="InputInstructions"] {
        display: none !important;
    }
    </style>
""", unsafe_allow_html=True)

# Injected directly into the parent window's head to prevent 'Enter' keypress from submitting forms
components.html("""
    <script>
    const parentDoc = window.parent.document;
    if (!parentDoc.getElementById('prevent-enter-submit')) {
        const script = parentDoc.createElement('script');
        script.id = 'prevent-enter-submit';
        script.type = 'text/javascript';
        script.innerHTML = `
            document.addEventListener('keydown', function(e) {
                if ((e.key === 'Enter' || e.keyCode === 13) && e.target.tagName === 'INPUT') {
                    e.preventDefault();
                    e.stopPropagation();
                    e.stopImmediatePropagation();
                }
            }, true);
        `;
        parentDoc.head.appendChild(script);
    }
    </script>
""", height=0, width=0)

# ==========================================
# TEMPLATE CONFIGURATION
# ==========================================
# Google Doc URL provided for YWL - Client Appointment Letter
YWL_TEMPLATE_URL = "https://docs.google.com/document/d/1LnriO5OPwb94aLdJh7tMsSnMpwnH5qIA/"

# ==========================================
# HELPER FUNCTIONS
# ==========================================
def extract_doc_id(url: str) -> str:
    """Extract Google Doc ID from full URL."""
    match = re.search(r"/d/([a-zA-Z0-9-_]+)", url)
    return match.group(1) if match else ""

def fetch_google_doc_bytes(url: str) -> bytes:
    """Download Google Doc directly as DOCX bytes."""
    doc_id = extract_doc_id(url)
    if not doc_id:
        raise ValueError(f"Invalid Google Doc URL provided: {url}")
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
    response = requests.get(export_url)
    response.raise_for_status()
    return response.content

def replace_placeholders_in_paragraph(paragraph, replacements: dict):
    """Replace placeholder keys in paragraph while preserving text runs and formatting."""
    for key, value in replacements.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)

def process_docx_bytes(file_bytes: bytes, replacements: dict) -> bytes:
    """Load DOCX bytes, replace placeholders globally, and return modified DOCX bytes."""
    doc_io = io.BytesIO(file_bytes)
    doc = Document(doc_io)

    # 1. Process standard paragraphs
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, replacements)

    # 2. Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, replacements)
                    
    # 3. Process headers and footers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                replace_placeholders_in_paragraph(p, replacements)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_placeholders_in_paragraph(p, replacements)
        if section.footer:
            for p in section.footer.paragraphs:
                replace_placeholders_in_paragraph(p, replacements)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_placeholders_in_paragraph(p, replacements)

    output_io = io.BytesIO()
    doc.save(output_io)
    return output_io.getvalue()

def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF bytes using headless LibreOffice."""
    with tempfile.TemporaryDirectory() as tmpdir:
        input_docx_path = os.path.join(tmpdir, "input.docx")
        with open(input_docx_path, "wb") as f:
            f.write(docx_bytes)
        
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", tmpdir,
            input_docx_path
        ]
        
        try:
            subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
            output_pdf_path = os.path.join(tmpdir, "input.pdf")
            if os.path.exists(output_pdf_path):
                with open(output_pdf_path, "rb") as f:
                    return f.read()
        except Exception as e:
            st.warning(f"PDF conversion warning: {e}")
    return None

# ==========================================
# USER INTERFACE
# ==========================================
st.title("YWL Mandate Agreement Generator")
st.markdown("Automated generation of the **Exclusive Financial Consultation & Mandate Agreement**.")

# Main Input Form
with st.form("doc_generation_form"):
    st.subheader("1. Client Details")
    col1, col2 = st.columns(2)
    
    with col1:
        client_name = st.text_input("Client Name", value="")
        client_nric = st.text_input("NRIC / Passport / Reg. No.", value="")
        client_email = st.text_input("Client Email", value="")
        
    with col2:
        client_contact = st.text_input("Contact Number", value="")
        agreement_date = st.date_input("Agreement Date", value=datetime.date.today())
        
    client_address = st.text_area("Correspondence Address", value="")

    st.subheader("2. Advisor Information")
    w_col1, w_col2 = st.columns(2)
    with w_col1:
        advisor_name = st.text_input("Advisor / Witness Name", value="")
    with w_col2:
        advisor_nric = st.text_input("Advisor / Witness NRIC", value="")

    submit_button = st.form_submit_button("Generate Agreement (PDF)")

# ==========================================
# PROCESSING
# ==========================================
if submit_button:
    with st.spinner("Fetching template and generating PDF..."):
        
        # --- Text Formatting Rules ---
        formatted_client_name = client_name.strip().upper()
        formatted_client_nric = client_nric.strip().upper()
        formatted_client_address = client_address.strip().title() if client_address.strip() else ""
        formatted_advisor_name = advisor_name.strip().upper()
        formatted_advisor_nric = advisor_nric.strip().upper()
        
        # Format date as e.g., "12 August 2026"
        formatted_agreement_date = agreement_date.strftime("%d %B %Y") if agreement_date else ""

        raw_replacements = {
            "<<CLIENT_NAME>>": formatted_client_name,
            "<<CLIENT_NRIC>>": formatted_client_nric,
            "<<CLIENT_EMAIL>>": client_email,
            "<<CLIENT_CONTACT>>": client_contact,
            "<<CLIENT_ADDRESS>>": formatted_client_address,
            "<<DATE>>": formatted_agreement_date,
            "<<ADVISOR_NAME>>": formatted_advisor_name,
            "<<ADVISOR_NRIC>>": formatted_advisor_nric,
        }

        # Map empty inputs to spaces so placeholders vanish cleanly if left empty
        replacements = {k: (v.strip() if v and str(v).strip() else "  ") for k, v in raw_replacements.items()}

        clean_file_client_name = client_name.strip().upper() if client_name.strip() else "CLIENT"
        clean_file_date = agreement_date.strftime("%Y%m%d") if agreement_date else ""
        
        file_name_base = f"[{clean_file_client_name}] YWL Mandate Agreement {clean_file_date}"

        try:
            # 1. Fetch Google Doc template
            doc_bytes_raw = fetch_google_doc_bytes(YWL_TEMPLATE_URL)

            # 2. Populate placeholders
            processed_docx = process_docx_bytes(doc_bytes_raw, replacements)

            # 3. Convert to PDF
            final_pdf = convert_docx_to_pdf(processed_docx)

            if final_pdf:
                st.success("✅ Document generated successfully!")
                
                st.download_button(
                    label="📥 Download Agreement (PDF)",
                    data=final_pdf,
                    file_name=f"{file_name_base}.pdf",
                    mime="application/pdf",
                    type="primary"
                )
            else:
                st.error("Failed to generate PDF. Please check server dependencies (LibreOffice).")

        except Exception as e:
            st.error(f"Error processing document: {e}")
